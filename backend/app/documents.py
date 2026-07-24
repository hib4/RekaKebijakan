from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import fitz
from docx import Document


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


@dataclass(frozen=True)
class TextSegment:
    text: str
    char_start: int
    char_end: int
    page: int | None = None
    paragraph: int | None = None
    line: int | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    segments: tuple[TextSegment, ...]


class ExtractedText(str):
    def __new__(cls, document: ExtractedDocument):
        value = super().__new__(cls, document.text)
        value.extraction = document
        return value


def _assemble(values: list[tuple[str, dict[str, int | None]]], max_chars: int | None = None) -> ExtractedDocument:
    text_parts: list[str] = []
    segments: list[TextSegment] = []
    cursor = 0
    for raw, locator in values:
        value = re.sub(r"\s+", " ", raw).strip()
        if not value:
            continue
        if text_parts:
            cursor += 1
        if max_chars is not None and cursor + len(value) > max_chars:
            raise ValueError("Teks hasil ekstraksi terlalu besar")
        start = cursor
        text_parts.append(value)
        cursor += len(value)
        segments.append(TextSegment(value, start, cursor, **locator))
    return ExtractedDocument(" ".join(text_parts), tuple(segments))


def extract_document(path: Path, max_pdf_pages: int = 200, max_chars: int = 2_000_000) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Format berkas tidak didukung: {suffix or 'tanpa ekstensi'}")
    values: list[tuple[str, dict[str, int | None]]] = []
    if suffix == ".pdf":
        with path.open("rb") as source:
            signature = source.read(5)
        if signature != b"%PDF-":
            raise ValueError("Tanda tangan berkas PDF tidak valid")
        with fitz.open(path) as document:
            if document.page_count > max_pdf_pages:
                raise ValueError("Jumlah halaman PDF melebihi batas")
            for page_index, page in enumerate(document, 1):
                payload = page.get_text("dict")
                line_number = 0
                for block_index, block in enumerate(payload.get("blocks", []), 1):
                    for line in block.get("lines", []):
                        line_number += 1
                        text = "".join(span.get("text", "") for span in line.get("spans", []))
                        values.append((text, {"page": page_index, "paragraph": block_index, "line": line_number}))
    elif suffix == ".docx":
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise ValueError("Kontainer DOCX tidak valid")
        except zipfile.BadZipFile as error:
            raise ValueError("Kontainer DOCX tidak valid") from error
        for paragraph_index, paragraph in enumerate(Document(path).paragraphs, 1):
            for line_index, line in enumerate(paragraph.text.splitlines() or [paragraph.text], 1):
                values.append((line, {"page": None, "paragraph": paragraph_index, "line": line_index}))
    else:
        paragraph = 0
        try:
            content = path.read_text(encoding="utf-8-sig", errors="strict")
        except UnicodeDecodeError as error:
            raise ValueError("Berkas teks harus berupa UTF-8 yang valid") from error
        if "\x00" in content:
            raise ValueError("Berkas teks biner tidak didukung")
        for line_index, line in enumerate(content.splitlines(), 1):
            if line.strip():
                paragraph += 1
            values.append((line, {"page": None, "paragraph": paragraph or None, "line": line_index}))
    extracted = _assemble(values, max_chars)
    if not extracted.text:
        raise ValueError("Dokumen tidak menghasilkan teks")
    return extracted


def extract_text(path: Path) -> str:
    return ExtractedText(extract_document(path))


def _chunk_metadata(extraction: ExtractedDocument | None, start: int, end: int) -> dict[str, Any]:
    if extraction is None:
        return {}
    matched = [segment for segment in extraction.segments if segment.char_start < end and segment.char_end > start]
    locators = [
        {
            "page": segment.page,
            "paragraph": segment.paragraph,
            "line": segment.line,
            "char_start": max(start, segment.char_start),
            "char_end": min(end, segment.char_end),
        }
        for segment in matched
    ]
    return {
        "pages": sorted({item["page"] for item in locators if item["page"] is not None}),
        "paragraphs": sorted({item["paragraph"] for item in locators if item["paragraph"] is not None}),
        "lines": locators,
    }


def chunk_text(
    document_id: str, text: str | ExtractedDocument, size: int = 1200, overlap: int = 150,
    max_chunks: int = 5000,
) -> list[dict]:
    if size < 100 or overlap < 0 or overlap >= size:
        raise ValueError("Konfigurasi chunk dokumen tidak valid")
    extraction = text if isinstance(text, ExtractedDocument) else getattr(text, "extraction", None)
    text = extraction.text if isinstance(extraction, ExtractedDocument) else str(text)
    chunks = []
    start = 0
    while start < len(text):
        if len(chunks) >= max_chunks:
            raise ValueError("Jumlah potongan dokumen melebihi batas")
        limit = min(len(text), start + size)
        end = limit
        if limit < len(text):
            boundary = text.rfind(" ", start + size // 2, limit)
            if boundary > start:
                end = boundary
        value = text[start:end].strip()
        if value:
            digest = hashlib.sha256(f"{document_id}:{len(chunks)}:{value}".encode()).hexdigest()
            chunks.append({
                "id": f"chunk_{digest[:16]}",
                "document_id": document_id,
                "ordinal": len(chunks),
                "text": value,
                "char_start": start,
                "char_end": end,
                "content_sha256": hashlib.sha256(value.encode()).hexdigest(),
                "metadata": _chunk_metadata(extraction, start, end),
            })
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks
